"""
Word report builder for DGA research reports.

Uses python-docx with manual XML for crisp, institutional-quality tables:
  - Every cell has visible black borders (not just the default 'Table Grid' style)
  - Header row has shaded background + bold white text
  - Numeric cells right-aligned
  - Consistent cell padding
  - Page headers with report title + date + page number
"""

from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


# ---------------------------------------------------------------------------
# Low-level XML helpers
# ---------------------------------------------------------------------------
def _set_cell_border(cell, color: str = "000000", size_eighth_pt: int = 8,
                     style: str = "single") -> None:
    """Force visible borders on all four sides of a cell (overrides Word defaults)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for side in ("top", "left", "bottom", "right"):
        elem = tcBorders.find(qn(f"w:{side}"))
        if elem is None:
            elem = OxmlElement(f"w:{side}")
            tcBorders.append(elem)
        elem.set(qn("w:val"), style)
        elem.set(qn("w:sz"), str(size_eighth_pt))
        elem.set(qn("w:color"), color)
        elem.set(qn("w:space"), "0")


def _set_cell_shading(cell, fill_hex: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)


def _set_cell_margins(cell, top: int = 80, bottom: int = 80,
                      left: int = 120, right: int = 120) -> None:
    """Cell padding in DXA (1440 = 1 inch). 80 ≈ 5.5pt top/bottom padding."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom),
                      ("left", left), ("right", right)):
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), str(val))
        m.set(qn("w:type"), "dxa")
        tcMar.append(m)
    existing = tcPr.find(qn("w:tcMar"))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(tcMar)


def _run_styled(paragraph, text: str, *, bold: bool = False, size: float = 10,
                color: RGBColor | None = None, font: str = "Calibri"):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = font
    if color is not None:
        run.font.color.rgb = color
    return run


# ---------------------------------------------------------------------------
# Numeric-cell detection
# ---------------------------------------------------------------------------
_NUMERIC_RE = re.compile(r"^-?\$?\s*[\(\-]?\d[\d,\.]*[\)%]?[A-Za-z]*$")


def _looks_numeric(text: str) -> bool:
    t = text.strip()
    if not t or t.upper() in {"N/A", "NA", "-", "—"}:
        return True  # right-align placeholders too
    # Accept things like "+12.3%", "(1,234.5)", "$5.6M", "$12.34", "1,234", "45.6ppt"
    if _NUMERIC_RE.match(t):
        return True
    if any(c.isdigit() for c in t) and len(t) < 20:
        return True
    return False


# ---------------------------------------------------------------------------
# Markdown parsing
# ---------------------------------------------------------------------------
_TABLE_ROW_RE = re.compile(r"^\s*\|(.+?)\|\s*$")
_SEP_ROW_RE = re.compile(r"^\s*\|?\s*(:?-{2,}:?\s*\|\s*)+:?-{2,}:?\s*\|?\s*$")


def _strip_md_inline(text: str) -> tuple[str, dict]:
    """Handle **bold** and *italic* and `code` minimally. Returns (plain_text, flags)."""
    flags = {"bold": False, "italic": False}
    t = text.strip()
    # Heuristic: if the WHOLE cell is wrapped in ** or *, flag it; else strip markers.
    if t.startswith("**") and t.endswith("**") and len(t) > 4:
        flags["bold"] = True
        t = t[2:-2]
    elif t.startswith("__") and t.endswith("__") and len(t) > 4:
        flags["bold"] = True
        t = t[2:-2]
    # Remove inline markers we don't render.
    t = re.sub(r"\*\*(.+?)\*\*", r"\1", t)
    t = re.sub(r"__(.+?)__", r"\1", t)
    t = re.sub(r"(?<!\w)\*(.+?)\*(?!\w)", r"\1", t)
    t = re.sub(r"`([^`]+)`", r"\1", t)
    return t, flags


def _split_row(line: str) -> list[str]:
    m = _TABLE_ROW_RE.match(line)
    if not m:
        return []
    inner = m.group(1)
    cells = [c.strip() for c in inner.split("|")]
    return cells


# ---------------------------------------------------------------------------
# Document building blocks
# ---------------------------------------------------------------------------
HEADER_FILL = "1F3A5F"   # DGA navy
HEADER_FONT_COLOR = RGBColor(0xFF, 0xFF, 0xFF)
ALT_ROW_FILL = "F2F5FA"  # very light navy tint


def _add_bordered_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    if not headers:
        return
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Even column widths using DXA. Content width on US Letter 1" margins = 9360 DXA
    # python-docx uses EMUs (914400/inch), so we set cell.width via Inches.
    total_content = 6.5  # inches — use 6.5" to leave nice margins
    # Weight the first column (label) wider than the numeric columns.
    first_col = min(2.2, max(1.6, total_content * 0.30))
    other = (total_content - first_col) / max(1, n_cols - 1)
    widths_in = [first_col] + [other] * (n_cols - 1)

    # Header row
    hdr = table.rows[0]
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        plain, _ = _strip_md_inline(text)
        _run_styled(p, plain, bold=True, size=10, color=HEADER_FONT_COLOR, font="Calibri")
        _set_cell_shading(cell, HEADER_FILL)
        _set_cell_border(cell)
        _set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.width = Inches(widths_in[i])

    # Data rows
    for r_idx, row_cells in enumerate(rows):
        row = table.rows[1 + r_idx]
        # Pad / truncate to header width.
        padded = (row_cells + [""] * n_cols)[:n_cols]
        for i, raw in enumerate(padded):
            cell = row.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            plain, flags = _strip_md_inline(raw)
            numeric = i > 0 and _looks_numeric(plain)
            p.alignment = (
                WD_ALIGN_PARAGRAPH.RIGHT if numeric else WD_ALIGN_PARAGRAPH.LEFT
            )
            _run_styled(p, plain, bold=flags["bold"], size=10, font="Calibri")
            _set_cell_border(cell)
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.width = Inches(widths_in[i])
            if r_idx % 2 == 1:
                _set_cell_shading(cell, ALT_ROW_FILL)

    # Tiny spacer after table
    doc.add_paragraph()


def _add_paragraph_with_inline_md(doc: Document, text: str) -> None:
    """Render a paragraph handling **bold**, *italic*, and bullet indicators."""
    leading_bullet = False
    stripped = text.lstrip()
    bullet_prefix = ""
    for marker in ("- ", "* ", "→ ", "• "):
        if stripped.startswith(marker):
            leading_bullet = True
            bullet_prefix = "• "
            stripped = stripped[len(marker):]
            break

    p = doc.add_paragraph()
    if leading_bullet:
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.15)
        _run_styled(p, bullet_prefix, size=11)

    # Split on bold/italic markers.
    # Order matters: handle bold first, then italic.
    tokens = re.split(r"(\*\*[^*]+\*\*|__[^_]+__|\*[^*\n]+\*|`[^`]+`)", stripped)
    for tok in tokens:
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            _run_styled(p, tok[2:-2], bold=True, size=11)
        elif tok.startswith("__") and tok.endswith("__"):
            _run_styled(p, tok[2:-2], bold=True, size=11)
        elif tok.startswith("*") and tok.endswith("*") and len(tok) > 2:
            r = _run_styled(p, tok[1:-1], size=11)
            r.italic = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = _run_styled(p, tok[1:-1], size=11, font="Consolas")
        else:
            _run_styled(p, tok, size=11)


# ---------------------------------------------------------------------------
# Header / cover
# ---------------------------------------------------------------------------
_BRANDING_DIR = Path(__file__).resolve().parent / "branding"


def _add_cover(doc: Document, ticker: str, entity: str, rating_hint: str = "",
               price: str | float | None = None, as_of: str | None = None) -> None:
    logo_path = _BRANDING_DIR / "dga_logo.png"
    if logo_path.exists():
        logo_para = doc.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_para.add_run().add_picture(str(logo_path), width=Inches(3.0))

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run_styled(title, "DGA CAPITAL RESEARCH", bold=True, size=20,
                color=RGBColor(0x1F, 0x3A, 0x5F))

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run_styled(sub, f"{entity} ({ticker})", bold=True, size=16)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line_parts = [as_of or datetime.now().strftime("%B %d, %Y")]
    if price not in (None, "", "N/A"):
        try:
            line_parts.append(f"Last Price: ${float(price):,.2f}")
        except Exception:
            line_parts.append(f"Last Price: {price}")
    if rating_hint:
        line_parts.append(f"Rating: {rating_hint}")
    _run_styled(meta, "  |  ".join(line_parts), size=11,
                color=RGBColor(0x55, 0x55, 0x55))

    # Horizontal rule via paragraph border.
    hr = doc.add_paragraph()
    pPr = hr._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F3A5F")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_heading(doc: Document, text: str, level: int) -> None:
    text = text.strip()
    if not text:
        return
    h = doc.add_heading(level=level)
    # Force Calibri + navy color for consistency.
    colors = {
        1: RGBColor(0x1F, 0x3A, 0x5F),
        2: RGBColor(0x2D, 0x55, 0x84),
        3: RGBColor(0x43, 0x6B, 0x9A),
    }
    sizes = {1: 16, 2: 13, 3: 12}
    _run_styled(h, text, bold=True, size=sizes.get(level, 12),
                color=colors.get(level, colors[3]), font="Calibri")


# ---------------------------------------------------------------------------
# Main entrypoint
# ---------------------------------------------------------------------------
def render_report(
    markdown_text: str,
    *,
    ticker: str,
    entity_name: str = "",
    output_path: str,
    price: str | float | None = None,
    rating_hint: str = "",
) -> str:
    """Render Grok's markdown report into a polished .docx."""
    doc = Document()

    # Global default font.
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Page margins (1 inch = default; explicit for safety).
    for section in doc.sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    _add_cover(doc, ticker, entity_name or ticker, rating_hint=rating_hint, price=price)

    lines = markdown_text.split("\n")
    i = 0
    n = len(lines)
    while i < n:
        raw = lines[i]
        stripped = raw.strip()

        # Headings
        if stripped.startswith("### "):
            _add_heading(doc, stripped[4:], 3)
            i += 1
            continue
        if stripped.startswith("## "):
            _add_heading(doc, stripped[3:], 2)
            i += 1
            continue
        if stripped.startswith("# "):
            _add_heading(doc, stripped[2:], 1)
            i += 1
            continue

        # Tables: a line starting with "|" AND the following line is a separator
        if (
            stripped.startswith("|")
            and i + 1 < n
            and _SEP_ROW_RE.match(lines[i + 1])
        ):
            header_cells = _split_row(stripped)
            i += 2  # skip header + separator
            rows: list[list[str]] = []
            while i < n and _TABLE_ROW_RE.match(lines[i]):
                rows.append(_split_row(lines[i]))
                i += 1
            _add_bordered_table(doc, header_cells, rows)
            continue

        # Blank line
        if not stripped:
            i += 1
            continue

        # Default: render paragraph with inline markdown
        _add_paragraph_with_inline_md(doc, stripped)
        i += 1

    doc.save(output_path)
    return output_path
